import os
import time
import shutil

def monitor_folder_and_trigger_reporting(folder_directory, completed_directory):
    """
    Continuously monitors a specified folder for .txt files (excluding 'EXAMPLE.txt').
    When a .txt file is found, it reads the first line (expected to be a directory path)
    and passes that path to the reporting function process_reporting(). 
    Once processed, the .txt file is moved to the 'completed' directory.
    """

    # Ensure the 'completed' directory exists; if not, create it
    os.makedirs(completed_directory, exist_ok=True)

    while True:
        # List all .txt files in the folder, excluding "EXAMPLE.txt"
        txt_files = [
            file 
            for file in os.listdir(folder_directory) 
            if file.endswith('.txt') and file != "EXAMPLE.txt"
        ]
        
        if txt_files:
            for txt_file in txt_files:
                try:
                    txt_file_path = os.path.join(folder_directory, txt_file)
                    
                    # Read the path from the .txt file
                    with open(txt_file_path, 'r') as f:
                        walk_path = f.readline().strip()
                        if not os.path.exists(walk_path):
                            print(f"Invalid path in {txt_file_path}: {walk_path}")
                            continue
                        print(f"Processing walk_path from {txt_file_path}: {walk_path}")
                    
                    # Call the reporting logic
                    process_reporting(walk_path)

                    # Move the processed .txt file to the 'completed' folder
                    completed_path = os.path.join(completed_directory, txt_file)
                    shutil.move(txt_file_path, completed_path)
                    print(f"Processed and moved: {txt_file_path} to {completed_path}")
                except Exception as e:
                    print(f"An error occurred while processing {txt_file_path}: {e}")
        
        # Wait for 1 minute before scanning again
        time.sleep(60)

def process_reporting(walk_path):
    """
    Generates reports (Word documents) based on coverage and variant data found in Excel/CSV files 
    located in a specified directory (walk_path). 

    This function:
    1. Reads and processes coverage data from a CSV file.
    2. Reads variant information from an Excel file.
    3. Sorts and filters out relevant variants (Tier I and Tier II).
    4. Populates a Word template with the findings.
    5. Saves either a "Variant Report" or a "Normal" report Word document, depending on findings.
    """

    # --- Import necessary libraries for this function ---
    import pandas as pd
    from docx import Document
    from docx.shared import Pt
    import requests
    import sys
    import re
    from docx.enum.text import WD_COLOR_INDEX
    from decimal import Decimal, ROUND_HALF_UP

    # Debug statements marking the start of the reporting
    print("###########################_REPORT_GENERATION_START_#########################################")
    print(walk_path)
    print("###########################_REPORT_GENERATION_START_#########################################")

    # Initialize a list for reading paths (if needed for expansions)
    read_path = []

    # -----------------------------------------------------
    # 1. Identify coverage and variant files in the target directory
    # -----------------------------------------------------
    # Attempt to locate .csv and .xlsx files with coverage/region data 
    # by iterating through the files in walk_path.
    try:
        for root, dirs, files in os.walk(walk_path):
            print(root, dirs, files)
            for file in files:
                print(file)
                # If the file name starts with a digit and has coverage info
                if file.startswith('2'):
                    file_name = file.lower()
                    print(file_name)
                    if "coverageregion" in file_name or "l001_report_" in file_name:
                        if "xlsx" in file_name:
                            VAR_df = file_name
                        if "csv" in file_name:
                            COV_df = file_name
    except Exception as e:
        print(f"An error occurred while processing the file {file}: {e}")

    # Helper to detect disease keywords from file paths
    def find_keywords(text):
        if not isinstance(text, str):
            return ''
        keywords = ['MDS', 'MPN', 'AML']
        found = [keyword for keyword in keywords if keyword in text]
        return ', '.join(set(found))

    # Construct full paths for coverage and variant data
    cov_dir = os.path.join(walk_path, COV_df)
    var_dir = os.path.join(walk_path, VAR_df)

    # This helps to identify which disease site we might be dealing with
    disease_site = find_keywords(var_dir)

    # -----------------------------------------------------
    # 2. Load variant Excel file
    # -----------------------------------------------------
    try:
        df_temp = pd.read_excel(var_dir)
        specific_value = '#Chr:ChrPos'
        # Find the index of the header
        header_index_df = df_temp[df_temp.iloc[:, 0] == specific_value].index[0] + 1
        # Load the actual DataFrame with the correct header
        df = pd.read_excel(var_dir, sheet_name='sorted', header=header_index_df)
    except:
        # If it fails to detect a header automatically, fallback
        df = pd.read_excel(var_dir, sheet_name='sorted', header=9)

    # -----------------------------------------------------
    # 3. Load Word templates (reports) and gene reference data
    #    (Paths replaced with placeholders for demonstration)
    # -----------------------------------------------------
    doc = Document(r'[PATH_TO_REPORT_TEMPLATE]\result_template.docx')
    doc_normal = Document(r'[PATH_TO_REPORT_TEMPLATE]\result_template_normal.docx')
    gene_info_df = pd.read_excel(
        r'[PATH_TO_GENE_STATEMENTS]\Gene_Statements_V2.xlsx', 
        sheet_name="Gene Statements"
    )
    previous_variants_df = pd.read_excel(
        r'[PATH_TO_GENE_STATEMENTS]\Gene_Statements_V2.xlsx', 
        sheet_name="Variants"
    )
    # Fill NaN in 'Functional_Evidence_Statement' so we can safely use it later
    previous_variants_df['Functional_Evidence_Statement'] = previous_variants_df['Functional_Evidence_Statement'].fillna('')

    # -----------------------------------------------------
    # 4. Load coverage CSV file
    # -----------------------------------------------------
    try:
        data_df_temp = pd.read_csv(cov_dir)
        specific_value = '#Transcript'
        header_index = data_df_temp[data_df_temp.iloc[:, 0] == specific_value].index[0] + 1
        data_df = pd.read_csv(cov_dir, header=header_index)
    except:
        # Fallback if standard approach fails
        data_df_temp = pd.read_csv(
            cov_dir, 
            names=[f'col_{i}' for i in range(16)]  # Example placeholder column names
        )
        specific_value = '#Transcript'
        header_index = data_df_temp[data_df_temp.iloc[:, 0] == specific_value].index[0]
        data_df = pd.read_csv(cov_dir, header=header_index)

    # -----------------------------------------------------
    # 5. Functions for variant classification and text formatting
    # -----------------------------------------------------

    # Mapping for amino acid 3-letter to 1-letter conversions
    amino_acid_mapping = {
        'Ala': 'A', 'Arg': 'R', 'Asn': 'N', 'Asp': 'D', 'Cys': 'C',
        'Glu': 'E', 'Gln': 'Q', 'Gly': 'G', 'His': 'H', 'Ile': 'I',
        'Leu': 'L', 'Lys': 'K', 'Met': 'M', 'Phe': 'F', 'Pro': 'P',
        'Ser': 'S', 'Thr': 'T', 'Trp': 'W', 'Tyr': 'Y', 'Val': 'V',
        'Ter': '*'
    }

    def convert_protein_notation(protein):
        """
        Converts HGVS protein notation (3-letter codons) to a 1-letter notation
        and properly formats frameshift nomenclature.
        """
        if not isinstance(protein, str):
            return "p.?"
        pattern = r'p\.([A-Za-z]{3})(\d+)([A-Za-z]{3}|Ter)?(fsTer\d+)?'
        match = re.match(pattern, protein)
        if not match:
            return protein
        
        orig_aa, position, change_aa, fs = match.groups()
        converted = [f'p.{amino_acid_mapping.get(orig_aa.capitalize(), orig_aa)}{position}']
        if change_aa:
            converted.append(amino_acid_mapping.get(change_aa.capitalize(), change_aa))
        if fs:
            fs = fs.replace('Ter', '*')
            converted.append(fs)
        return ''.join(converted)

    def determine_tier(row):
        """
        Determine Tier number based on 'Classification' or 'Pathogenicity' columns.
        """
        variant_comment = row['Classification']
        pathogenicity = row['Pathogenicity']

        if pd.notna(variant_comment):
            if 'Tier II' in variant_comment:
                return 2
            elif 'Tier I' in variant_comment:
                return 1

        if 'Tier II' in pathogenicity:
            return 2
        elif 'Tier I' in pathogenicity:
            return 1

        return None

    def apply_font_style(paragraph, font_name='Arial', font_size=10):
        """
        Applies a consistent font style to a paragraph.
        """
        paragraph.style.font.name = font_name
        paragraph.style.font.size = Pt(font_size)

    def get_most_recent_functional_evidence(gene, protein_change, disease_site, previous_variants_df):
        """
        Looks up the functional evidence statement (if any) for a given gene and protein change 
        in a reference 'previous_variants_df' DataFrame. 
        Considers disease site to narrow results.
        """
        new_variant = False
        matching_rows = previous_variants_df[
            (previous_variants_df['Gene'] == gene) & 
            (previous_variants_df['Protein'].str.contains(protein_change, na=False)) &
            (previous_variants_df['Disease_Site'].str.contains(disease_site, na=False))
        ]
        
        if not matching_rows.empty:
            most_recent_row = matching_rows.loc[matching_rows['Index'].idxmax()]
            return most_recent_row['Functional_Evidence_Statement'], new_variant
        else:
            new_variant = True
            return "", new_variant

    # -----------------------------------------------------
    # 6. Filtering out Tier III variants and focusing on Tier I/II
    # -----------------------------------------------------
    df['Classification'] = df['Classification'].fillna('').astype(str)
    df['Pathogenicity'] = df['Pathogenicity'].fillna('').astype(str)

    exclude_tier_iii_filter = ~(
        df['Classification'].str.contains('Tier III', na=False) |
        df['Pathogenicity'].str.contains('Tier III', na=False)
    )
    include_tier_i_ii_filter = (
        df['Classification'].str.contains('^Tier I$|^Tier II$', regex=True, na=False) |
        df['Pathogenicity'].isin(['Tier I', 'Tier II'])
    )
    final_filter = exclude_tier_iii_filter & include_tier_i_ii_filter
    df_filtered = df[final_filter]

    try:
        df_filtered['TierNumber'] = df_filtered.apply(determine_tier, axis=1)
        df_filtered_sorted = df_filtered.sort_values(by=['TierNumber', 'Gene'])

        # Compile gene names for the [GENE_LIST] placeholder
        gene_names = df_filtered_sorted['Gene'].tolist()
        if len(gene_names) > 1:
            if len(gene_names) == 2:
                gene_list_sentence = ' and '.join(gene_names)
            else:
                gene_list_sentence = ', '.join(gene_names[:-1]) + f" and {gene_names[-1]}"
        else:
            gene_list_sentence = gene_names[0]

        # -----------------------------------------------------
        # 7. Populate the template table with the filtered variants
        # -----------------------------------------------------
        table = doc.tables[0]
        start_row = 1  # Start filling data from the second row (row 0 = header)
        filled_rows = 0

        # Ensure the required columns exist
        if not {'HGVSCodingTranscript', 'HGVSTranslationProtein'}.issubset(df_filtered_sorted.columns):
            # Split columns from the existing combined columns if needed
            df_filtered_sorted[['HGVSCodingTranscript', 'HGVSCoding']] = df_filtered_sorted['HGVSCoding'].str.split(":", expand=True)
            df_filtered_sorted[['HGVSTranslationProtein', 'HGVSProtein']] = df_filtered_sorted['HGVSProtein'].str.split(":", expand=True)

        for index, row in df_filtered_sorted.iterrows():
            row_index = start_row + filled_rows
            if row_index >= len(table.rows):
                break  # No more rows in the template table

            table_row = table.rows[row_index]
            # Column 0: Gene and Transcript
            gene_paragraph = table_row.cells[0].paragraphs[0]
            gene_transcript_text = f"{row['Gene']}\n({row['HGVSCodingTranscript']})\n"
            gene_paragraph.text = gene_transcript_text
            apply_font_style(gene_paragraph)

            # Column 1: HGVS Coding & Protein
            hgvs_paragraph = table_row.cells[1].paragraphs[0]
            hgvs_paragraph.text = f"{row['HGVSCoding']}"
            if pd.notna(row['HGVSProtein']):
                protein_data = row['HGVSProtein'].strip()
                if protein_data.startswith('p.'):
                    protein_data = protein_data[2:]
            else:
                protein_data = ""
            protein_data = f"p.({protein_data})" if protein_data else "p.()"
            hgvs_paragraph.text += f"\n{protein_data}"
            apply_font_style(hgvs_paragraph)

            # Column 2: Coverage
            coverage_paragraph = table_row.cells[2].paragraphs[0]
            coverage_paragraph.text = f"{row['Coverage']:,.0f}x"
            apply_font_style(coverage_paragraph)

            # Column 3: Variant Frequency
            frequency_paragraph = table_row.cells[3].paragraphs[0]
            variant_frequency = row['VariantFrequency'] * 100
            frequency_text = f"{variant_frequency:.1f}%" if variant_frequency % 1 else f"{int(variant_frequency)}%"
            frequency_paragraph.text = frequency_text
            apply_font_style(frequency_paragraph)

            # Column 4: Classification
            classification_paragraph = table_row.cells[4].paragraphs[0]
            classification_text = row['Classification'] if row['Classification'].strip() else row['Pathogenicity']
            classification_paragraph.text = classification_text
            apply_font_style(classification_paragraph)

            filled_rows += 1

        # Remove any empty rows in the template table after data
        for _ in range(start_row + filled_rows, len(table.rows)):
            table._element.remove(table.rows[start_row + filled_rows]._tr)

    except:
        # If no Tier I/II data found
        df_filtered_sorted = pd.DataFrame()

    # -----------------------------------------------------
    # 8. Compute and format coverage data
    # -----------------------------------------------------
    average_depth_of_coverage = data_df['AverageCoverage'].mean()
    formatted_average_depth_of_coverage = f"{average_depth_of_coverage:,.0f}x"

    # Exclude certain regions with underscores or special naming
    data_df['RegionName'] = data_df['RegionName'].str.upper()
    regions_not_reported = data_df[
        (data_df['RegionName'].str.count('_') > 1) |
        (data_df['RegionName'] == "ZRSR2_ENSE00001919416")
    ]
    data_df.drop(index=regions_not_reported.index, errors='ignore', inplace=True)

    # Create columns for gene/exon from RegionName
    data_df[['GeneName', 'ExonNumber']] = data_df['RegionName'].str.split('_EX', expand=True)
    data_df['ExonNumber_INT'] = pd.to_numeric(
        data_df['ExonNumber'].str.extract(r'(\d+)')[0],
        errors='coerce'
    ).astype(pd.Int32Dtype())
    data_df['Gene_Exon'] = data_df['#Transcript'] + ' (' + data_df['GeneName'] + ') exon ' + data_df['ExonNumber'].astype(str)

    # Identify coverage <250x
    low_cov_genes = data_df[data_df['MinimumCoverage'] < 250][
        ['RegionName', 'Gene_Exon', 'GeneName', 'ExonNumber', 'ExonNumber_INT', '#Transcript']
    ]

    low_cov_statement = False
    if len(low_cov_genes) > 25:
        low_cov_statement = True

    low_cov_genes = low_cov_genes.sort_values(by=['GeneName', 'ExonNumber_INT'])

    # Build a dictionary of gene-exon data
    gene_exon_dict = {}
    for _, row in low_cov_genes.iterrows():
        gene = row['#Transcript'] + " (" + row['GeneName'] + ")"
        exon = row['ExonNumber']
        if gene not in gene_exon_dict:
            gene_exon_dict[gene] = [exon]
        else:
            gene_exon_dict[gene].append(exon)

    formatted_genes_list = []
    for gene, exons in gene_exon_dict.items():
        if len(exons) > 1:
            formatted_exons = f'exons {", ".join(map(str, exons))}'
        else:
            formatted_exons = f'exon {exons[0]}'
        formatted_genes_list.append(f'{gene} {formatted_exons}')
    genes_list_low_cov = '; '.join(formatted_genes_list)

    # Next, compress consecutive exons into ranges (e.g., exon 3-5)
    import re
    gene_exon_pairs = genes_list_low_cov.split(';')
    grouped_exons = {}

    for pair in gene_exon_pairs:
        match = re.match(r'(.+?)\s+exons?\s+(.+)', pair)
        if match:
            gene = match.group(1)
            exon_str = match.group(2)
            exon_numbers = [int(e) for e in re.findall(r'\d+', exon_str)]
            current_group = [exon_numbers[0]]
            prev_exon = exon_numbers[0]
            for exon in exon_numbers[1:]:
                if exon == prev_exon + 1:
                    current_group.append(exon)
                else:
                    if len(current_group) >= 3:
                        grouped_exons.setdefault(gene, []).append(current_group)
                    else:
                        for individual_exon in current_group:
                            grouped_exons.setdefault(gene, []).append([individual_exon])
                    current_group = [exon]
                prev_exon = exon
            if len(current_group) >= 3:
                grouped_exons.setdefault(gene, []).append(current_group)
            else:
                for individual_exon in current_group:
                    grouped_exons.setdefault(gene, []).append([individual_exon])

    formatted_gene_list = []
    for gene, groups in grouped_exons.items():
        formatted_groups = []
        for group in groups:
            if len(group) > 1:
                formatted_group = f'{group[0]}-{group[-1]}'
            else:
                formatted_group = str(group[0])
            formatted_groups.append(formatted_group)
        if len(formatted_groups) > 1:
            formatted_gene_exons = ', '.join(formatted_groups[:-1]) + ' and ' + formatted_groups[-1]
        else:
            formatted_gene_exons = formatted_groups[0]
        formatted_gene_list.append(f'{gene} exon{"s" if len(groups) > 1 else ""} {formatted_gene_exons}')

    grouped_genes_list_low_cov = ';'.join(formatted_gene_list) + '.'
    formatted_avg_cov = f"{average_depth_of_coverage:,.0f}x"

    # Simple placeholder replacement in paragraphs
    def replace_placeholder(paragraph, placeholder, value, font_name='Arial', font_size=10, italic=False):
        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder in full_text:
            paragraph.clear()
            new_run = paragraph.add_run(full_text.replace(placeholder, str(value)))
            new_run.font.name = font_name
            new_run.font.size = Pt(font_size)
            new_run.italic = italic

    def replace_placeholder_with_highlight(document, placeholder, interpretation_text, is_new_variant):
        """
        Example function if you wanted to highlight placeholders for new variants
        in the Word document. Not fully used in this script, but provided as a reference.
        """
        for paragraph in document.paragraphs:
            if placeholder in paragraph.text:
                paragraph.clear()
                p_run = paragraph.add_run(interpretation_text)
                if is_new_variant:
                    p_run.font.highlight_color = WD_COLOR_INDEX.GREEN

    # Define the text for coverage statement
    if low_cov_statement:
        low_cov_text = (
            "Please note that poor sequence coverage was noted at several loci (refer to list in result section). "
            "If sequence variants were present at low allele frequencies in these regions then they may not have been detectable."
        )
    else:
        low_cov_text = ""

    # -----------------------------------------------------
    # 9. Handle the case where no Tier I/II variants found
    # -----------------------------------------------------
    if df_filtered_sorted.empty:
        # If no variants are reported, fill placeholders in the 'normal' template
        for paragraph in doc_normal.paragraphs:
            replace_placeholder(paragraph, "[AVG_COV]", formatted_avg_cov)
            replace_placeholder(paragraph, "[GENE_LIST_LOW_COV]", grouped_genes_list_low_cov)
            replace_placeholder(paragraph, "[LOW_COV_TEXT]", low_cov_text, italic=True)

        doc_normal.save(os.path.join(walk_path, "Variant_Report_normal.docx"))
    else:
        # -----------------------------------------------------
        # 10. If Tier I/II variants exist, fill in the main template
        # -----------------------------------------------------
        for paragraph in doc.paragraphs:
            replace_placeholder(paragraph, "[AVG_COV]", formatted_avg_cov)
            replace_placeholder(paragraph, "[GENE_LIST_LOW_COV]", grouped_genes_list_low_cov)
            replace_placeholder(paragraph, "[GENE_LIST]", gene_list_sentence, italic=True)
            replace_placeholder(paragraph, "[LOW_COV_TEXT]", low_cov_text, italic=True)

        # -----------------------------------------------------
        # 10a. Build the variant interpretation block
        # -----------------------------------------------------

        # Identify the type of variant (e.g., frameshift, nonsense, missense)
        def identify_variant_type(row):
            protein = row['HGVSProtein']
            coding = row['HGVSCoding']
            splice = row['ExonNumber']

            # If there's an intronic splice-site notation in the coding field
            if '+' in str(coding) or '-' in str(coding) and 'intron' in str(splice):
                return 'Splice-site'

            if isinstance(protein, str):
                if 'fsTer' in protein:
                    return 'Frameshift'
                elif 'delins' in protein:
                    return 'Insertion-Deletion'
                elif 'ins' in protein:
                    return 'Insertion'
                elif 'del' in protein:
                    return 'Deletion'
                elif 'dup' in protein:
                    return "Duplication"
                elif 'Ter' in protein:
                    return 'Nonsense'
                elif 'p.' in protein and protein.count('.') == 1 and protein.count('>') == 0:
                    return 'Missense'
                else:
                    return 'Other'

        df_filtered_sorted['VariantType'] = df_filtered_sorted.apply(identify_variant_type, axis=1)

        # Basic templates for different variant types
        templates = {
            'Frameshift': (
                "This sample is positive for a [NUMBER_OF_BP] bp [DELINS] in the [GENE_NAME] gene "
                "(HGVS nomenclature: [TRANSCRIPT]([GENE_NAME]):[CODING], [PROTEIN]), which was present "
                "in approximately [VAF]% of the sequenced fragments. This variant results in a frameshift "
                "of the coding sequence and is predicted to disrupt protein function by either premature "
                "protein truncation or nonsense mediated RNA decay."
            ),
            'Nonsense': (
                "This sample is positive for a nonsense variant in the [GENE_NAME] gene "
                "(HGVS nomenclature: [TRANSCRIPT]([GENE_NAME]):[CODING], [PROTEIN]) which was present in "
                "approximately [VAF]% of the sequenced fragments. This variant results in a premature stop "
                "codon and is predicted to disrupt protein function."
            ),
            'Missense': (
                "This sample is positive for a missense variant in the [GENE_NAME] gene "
                "(HGVS nomenclature: [TRANSCRIPT]([GENE_NAME]):[CODING], [PROTEIN]), which was present "
                "in approximately [VAF]% of the sequenced fragments. This variant substitutes [AA_REF] by "
                "[AA_VAR] at position [CODON], potentially affecting protein function (PMID: ...)."
            ),
            'Splice-site': (
                "This sample is positive for a [NUMBER_OF_BP] [DELINS] variant at the consensus splice site "
                "in the [GENE_NAME] gene (HGVS nomenclature: [TRANSCRIPT]([GENE_NAME]):[CODING], [PROTEIN]) "
                "which was present in approximately [VAF]% of the sequenced fragments. This variant is predicted "
                "to disrupt normal splicing and potentially protein function."
            ),
            'Insertion-Deletion': (
                "This sample is positive for an in-frame deletion-insertion variant in the [GENE_NAME] gene "
                "(HGVS nomenclature: [TRANSCRIPT]([GENE_NAME]):[CODING], [PROTEIN]) which was present "
                "in approximately [VAF]% of the sequenced fragments."
            ),
            'Deletion': (
                "This sample is positive for an in-frame deletion variant in the [GENE_NAME] gene "
                "(HGVS nomenclature: [TRANSCRIPT]([GENE_NAME]):[CODING], [PROTEIN]) which was present in "
                "approximately [VAF]% of the sequenced fragments."
            ),
            'Insertion': (
                "This sample is positive for an in-frame insertion variant in the [GENE_NAME] gene "
                "(HGVS nomenclature: [TRANSCRIPT]([GENE_NAME]):[CODING], [PROTEIN]) which was present "
                "in approximately [VAF]% of the sequenced fragments."
            ),
            'Duplication': (
                "This sample is positive for a [NUMBER_OF_BP] bp duplication in the [GENE_NAME] gene "
                "(HGVS nomenclature: [TRANSCRIPT]([GENE_NAME]):[CODING], [PROTEIN]) which was present "
                "in approximately [VAF]% of the sequenced fragments."
            )
        }

        def parse_hgvs_notation(hgvs_str):
            """
            Attempts to parse standard c.### notation for del/ins/dup to extract the 
            number of affected base pairs.
            """
            pattern = re.compile(
                r"c\."
                r"(\d+)"
                r"(_(\d+))?"
                r"(del|dup|ins)"
                r"([ACGT]+)?"
            )
            match = pattern.match(hgvs_str)
            if not match:
                return None
            start_position = int(match.group(1))
            end_position = int(match.group(3)) if match.group(3) else start_position
            mutation_type = match.group(4)
            nucleotides = match.group(5)
            return {
                'start_position': start_position,
                'end_position': end_position,
                'mutation_type': mutation_type,
                'nucleotides': nucleotides
            }

        def parse_hgvs_notation_splice(hgvs_str):
            """
            Specialized parsing for splice site HGVS notations that may include offsets like c.123-2.
            """
            pattern = re.compile(
                r"c\."
                r"(\d+-?\d*)"
                r"(_(\d+-?\d*))?"
                r"(del|dup|ins)"
                r"([ACGT]+)?"
            )
            match = pattern.match(hgvs_str)
            if not match:
                return None
            
            # In a real scenario, you would parse the numeric portion carefully.
            # Below is a simple stand-in for demonstration:
            start_position = match.group(1)
            end_position = match.group(3) if match.group(3) else start_position
            mutation_type = match.group(4)
            nucleotides = match.group(5)
            # ...
            return {
                'start_position': start_position,
                'end_position': end_position,
                'mutation_type': mutation_type,
                'nucleotides': nucleotides
            }

        def fill_template(row):
            """
            Fills the appropriate template based on the VariantType and placeholders in the 
            variant description.
            """
            variant_type = row['VariantType']
            template = templates.get(variant_type, "")
            if not template:
                return ""

            template = template.replace("[GENE_NAME]", row['Gene'])
            template = template.replace("[TRANSCRIPT]", row['HGVSCodingTranscript'])
            template = template.replace("[CODING]", row['HGVSCoding'])

            if isinstance(row['HGVSProtein'], str):
                protein_data = "p.(" + row['HGVSProtein'].split('p.')[-1] + ")"
            else:
                protein_data = "p.()"
            template = template.replace("[PROTEIN]", protein_data)

            ROUNDED_VAF = Decimal(row['VariantFrequency'] * 100).quantize(Decimal('1'), ROUND_HALF_UP)
            template = template.replace("[VAF]", str(ROUNDED_VAF))

            # Missense placeholders
            if variant_type == 'Missense':
                aa_parts = row['HGVSProtein'].split('.')
                if len(aa_parts) > 1 and len(aa_parts[1]) > 3:
                    aa_ref = aa_parts[1][:3]
                    aa_var = aa_parts[1][-3:]
                    codon = aa_parts[1][3:-3]
                    template = template.replace("[AA_REF]", aa_ref)
                    template = template.replace("[AA_VAR]", aa_var)
                    template = template.replace("[CODON]", codon)

            # Handle frameshift/indel calculations, splice, etc.
            coding_sequence = row['HGVSCoding']
            if variant_type in ['Frameshift','Deletion',"Insertion","Duplication"]:
                if "del" in coding_sequence or "dup" in coding_sequence or "ins" in coding_sequence:
                    match = parse_hgvs_notation(coding_sequence)
                    if match:
                        start = match['start_position']
                        end = match['end_position']
                        length = end - start + 1
                        if match['mutation_type'] == "del":
                            template = template.replace("[NUMBER_OF_BP]", str(length))
                            template = template.replace("[DELINS]", "deletion")
                        elif match['mutation_type'] == "ins":
                            # length = length of 'nucleotides'
                            if match['nucleotides']:
                                template = template.replace("[NUMBER_OF_BP]", str(len(match['nucleotides'])))
                            template = template.replace("[DELINS]", "insertion")
                        elif match['mutation_type'] == "dup":
                            template = template.replace("[NUMBER_OF_BP]", str(length))
                            template = template.replace("[DELINS]", "duplication")

            if variant_type == "Splice-site":
                if "del" in coding_sequence or "dup" in coding_sequence or "ins" in coding_sequence:
                    match = parse_hgvs_notation_splice(coding_sequence)
                    if match:
                        # For demonstration, you might do more advanced calculations here
                        template = template.replace("[NUMBER_OF_BP]", "[N_BPS]")
                        template = template.replace("[DELINS]", match['mutation_type'])
                else:
                    template = template.replace("[NUMBER_OF_BP]","")
                    template = template.replace("[DELINS]", "sequence")

            return template

        df_filtered_sorted['VariantDescription'] = df_filtered_sorted.apply(fill_template, axis=1)

        # -----------------------------------------------------
        # 10b. Query COSMIC data (optional external API usage)
        # -----------------------------------------------------
        def get_cosmic_data(variant, gene, coding):
            """
            Queries an external (COSMIC) database for variant data. 
            This function can be disabled or replaced as needed.
            """
            try:
                # Basic demonstration of building an API URL
                url = f'https://clinicaltables.nlm.nih.gov/api/cosmic/v4/search?terms={gene}+{coding}'
                response = requests.get(url)
                if response.status_code == 200:
                    data = response.json()
                    records = data[3]
                    if records:
                        return records[0]  # Return first record
                    else:
                        return [None] * 18
                else:
                    return [None] * 18
            except:
                return ["COSMIC_Error"] * 18

        cosmic_columns = [
            "AccessionNumber", "GeneCDS_Length", "GeneName", "HGNC_ID", 
            "MutationAA", "MutationCDS", "MutationDescription", 
            "MutationGenomePosition", "MutationStrand", "MutationID", 
            "LegacyMutationID", "GenomicMutationID", "Name", 
            "PrimaryHistology", "PrimarySite", "PubmedPMID", 
            "Site", "GRChVer"
        ]

        cosmic_df = pd.DataFrame(
            df_filtered_sorted.apply(
                lambda row: get_cosmic_data(row['VariantType'], row['Gene'], row['HGVSCoding']), 
                axis=1
            ).tolist(), 
            columns=cosmic_columns
        )

        # Convert HGVS Protein to simpler notation for matching
        df_filtered_sorted['ConvertedProtein'] = df_filtered_sorted['HGVSProtein'].apply(convert_protein_notation)

        merged_df = pd.merge(
            df_filtered_sorted,
            cosmic_df,
            left_on=['Gene', 'ConvertedProtein'],
            right_on=['GeneName', 'MutationAA'],
            how='left'
        )

        # Build a COSMIC statement
        def cosmic_statement(row):
            primary_site = row['PrimarySite']
            legacy_mutation_id = row['LegacyMutationID']
            pubmed_pmids = row['PubmedPMID']

            pmid_count = len(pubmed_pmids.split(';')) if pd.notna(pubmed_pmids) else 0

            if pd.isna(primary_site) or primary_site == "":
                return "This variant has not been previously catalogued through the COSMIC database in haematopoietic and lymphoid malignancies."

            elif 'haematopoietic_and_lymphoid_tissue' in primary_site:
                statement = (
                    f"This variant has been catalogued in the COSMIC database as a rare variant "
                    f"in haematopoietic and lymphoid malignancies (mutation ID {legacy_mutation_id})."
                )
                if pmid_count > 10:
                    statement += (
                        f" [CHECK IF HOTSPOT] This variant occurs at a hotspot position (mutation ID {legacy_mutation_id})."
                    )
                return statement
            else:
                return (
                    f"This variant has been catalogued in the COSMIC database but not in haematopoietic "
                    f"and lymphoid malignancies (mutation ID {legacy_mutation_id})."
                )

        merged_df['cosmic_statement'] = merged_df.apply(cosmic_statement, axis=1)

        # Helper to retrieve a gene-level statement
        def get_gene_statement(gene, gene_info_df):
            matched_gene = gene_info_df[gene_info_df['Gene'] == gene]
            if not matched_gene.empty:
                return matched_gene.iloc[0]['Statement']
            else:
                return "No specific gene statement available."

        merged_df['gene_level_statement_column'] = merged_df.apply(
            lambda row: get_gene_statement(row['Gene'], gene_info_df), 
            axis=1
        )

        # Flag if VAF is > 50% to add a statement
        merged_df['VAF_statement'] = merged_df['VariantFrequency'].apply(
            lambda vaf: (
                "Observation of this variant at an allele frequency > 50% suggests "
                "it may be present as a biallelic or hemizygous variant.\n"
                if vaf > 0.5 else ""
            )
        )

        # -----------------------------------------------------
        # 10c. Insert the [INTERPRETATION] text in the doc
        # -----------------------------------------------------
        interpretation_texts = []
        for index, row in merged_df.iterrows():
            gene_name = row['Gene']
            variant_description = row['VariantDescription']
            cosmic_stmt = row['cosmic_statement']
            gene_lvl_stmt = row['gene_level_statement_column']
            vaf_stmt = row['VAF_statement']

            # Construct final text sections
            protein_change = f"p.{row['HGVSProtein'].lstrip('p.')}" if pd.notna(row['HGVSProtein']) else "p.()"
            functional_evidence_statement, is_new_variant = get_most_recent_functional_evidence(
                gene_name, protein_change, disease_site, previous_variants_df
            )
            functional_evidence = functional_evidence_statement if functional_evidence_statement else ""

            # Build combined text
            gene_name_text = f"{gene_name} variant description: "
            variant_info_text = f"{variant_description} {cosmic_stmt}\n\n"
            gene_statement_text = f"{gene_lvl_stmt}\n\n{vaf_stmt}\n"

            interpretation_texts.append(
                (gene_name_text, variant_info_text, gene_statement_text, functional_evidence, is_new_variant)
            )

        # Replace [INTERPRETATION] placeholder
        for paragraph in doc.paragraphs:
            if "[INTERPRETATION]" in paragraph.text:
                paragraph.clear()
                for (
                    gene_name_text, 
                    variant_info_text, 
                    gene_statement_text, 
                    functional_evidence, 
                    is_new_variant
                ) in interpretation_texts:
                    # Underline gene name
                    gene_name_run = paragraph.add_run(gene_name_text)
                    gene_name_run.underline = True
                    
                    # Variant + cosmic statement
                    paragraph.add_run(variant_info_text)

                    # If there's a functional evidence, highlight it
                    if functional_evidence:
                        fe_run = paragraph.add_run(functional_evidence)
                        fe_run.font.highlight_color = WD_COLOR_INDEX.RED
                        paragraph.add_run("\n\n")

                    # Gene-level statement 
                    paragraph.add_run(gene_statement_text)

                    # Highlight gene name if it's a newly discovered variant
                    if is_new_variant:
                        gene_name_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        # -----------------------------------------------------
        # 10d. Final naming and saving of the Word document
        # -----------------------------------------------------
        variant_strings = []
        for index, row in df_filtered_sorted.iterrows():
            if isinstance(row['HGVSProtein'], str):
                variant_strings.append(f"{row['Gene']} p.{row['HGVSProtein'].lstrip('p.')}")
            else:
                variant_strings.append(f"{row['Gene']} p.()")

        combined_variants = ';'.join(variant_strings)
        final_file_name = f"_Variant Report_{combined_variants.replace(';', '_')}.docx"
        doc.save(os.path.join(walk_path, final_file_name))

    print("###########################_REPORT_GENERATION_COMPLETE_#########################################")


if __name__ == "__main__":
    # Example usage of the script. Update the paths below as needed.
    folder_directory = r"[PATH_TO_WATCHED_FOLDER]"  
    completed_directory = r"[PATH_TO_WATCHED_FOLDER]\completed"
    monitor_folder_and_trigger_reporting(folder_directory, completed_directory)
